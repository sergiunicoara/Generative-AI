"""Load PDF, DOCX, TXT, and MD files into Document objects."""

from __future__ import annotations

from pathlib import Path

import structlog

from graphrag.core.models import Document

log = structlog.get_logger(__name__)


def load_document(file_path: str | Path) -> Document:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _load_pdf(path)
    elif suffix == ".docx":
        text = _load_docx(path)
    elif suffix in (".txt", ".md"):
        text = path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    doc = Document(
        filename=path.name,
        source_path=str(path.resolve()),
        raw_text=text,
        metadata={"extension": suffix, "size_bytes": path.stat().st_size},
    )
    log.info("document_loader.loaded", filename=doc.filename, chars=len(text))
    return doc


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(para.text for para in doc.paragraphs)
