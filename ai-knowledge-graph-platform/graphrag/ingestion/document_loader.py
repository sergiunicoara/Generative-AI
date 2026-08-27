"""Load PDF, DOCX, TXT, Markdown, JSON, and Excel files into Documents."""

from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path
from html.parser import HTMLParser
from urllib.parse import urljoin

import structlog

from graphrag.core.content_hash import compute_content_hash
from graphrag.core.models import Document, StructuredTable
from graphrag.enterprise.models import DocumentLink

log = structlog.get_logger(__name__)


def load_document(file_path: str | Path) -> Document:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    suffix = path.suffix.lower()
    text = load_document_content(path.name, path.read_bytes())

    doc = Document(
        filename=path.name,
        source_path=str(path.resolve()),
        raw_text=text,
        metadata={"extension": suffix, "size_bytes": path.stat().st_size},
        # Computed here, at the single point every format converges to str,
        # so the whole pipeline downstream can compare "is this the same
        # document as last run?" without re-reading the file.
        content_hash=compute_content_hash(text),
    )
    doc.outbound_links = extract_document_links(path.name, path.read_bytes())
    tables = load_structured_tables(path.name, path.read_bytes(), document_id=doc.id, tenant=doc.tenant)
    if tables:
        doc.metadata["structured_tables"] = [table.model_dump(mode="json") for table in tables]
    log.info("document_loader.loaded", filename=doc.filename, chars=len(text))
    return doc


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(para.text for para in doc.paragraphs)


def load_document_content(filename: str, content: bytes) -> str:
    """Extract text from a local or remotely downloaded document payload."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader
        return "\n\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
    if suffix == ".docx":
        from docx import Document as DocxDocument
        return "\n".join(para.text for para in DocxDocument(BytesIO(content)).paragraphs)
    if suffix == ".json":
        return json.dumps(json.loads(content.decode("utf-8")), ensure_ascii=False, indent=2, sort_keys=True)
    if suffix == ".xlsx":
        from openpyxl import load_workbook
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in workbook.worksheets:
            rows.append(f"# Sheet: {sheet.title}")
            rows.extend(" | ".join("" if value is None else str(value) for value in row)
                        for row in sheet.iter_rows(values_only=True))
        return "\n".join(rows)
    if suffix in {".html", ".htm"}:
        parser = _HTMLTextAndLinks()
        parser.feed(content.decode("utf-8", errors="replace"))
        parser.close()
        return "\n".join(part for part in parser.text if part).strip()
    if suffix in {".txt", ".md", ".csv"}:
        return content.decode("utf-8")
    raise ValueError(f"Unsupported file type: {suffix}")


class _HTMLTextAndLinks(HTMLParser):
    """Small dependency-free HTML reader for visible text and explicit anchors."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.text: list[str] = []
        self.links: list[tuple[str, str, int]] = []
        self._ignored_depth = 0
        self._current_href = ""
        self._current_anchor: list[str] = []
        self._link_index = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._ignored_depth += 1
        if tag == "a" and not self._ignored_depth:
            self._current_href = dict(attrs).get("href") or ""
            self._current_anchor = []

    def handle_endtag(self, tag: str) -> None:
        if tag == "a" and self._current_href:
            self._link_index += 1
            self.links.append((self._current_href, " ".join(self._current_anchor).strip(), self._link_index))
            self._current_href = ""
            self._current_anchor = []
        if tag in {"script", "style", "noscript"} and self._ignored_depth:
            self._ignored_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._ignored_depth:
            return
        value = " ".join(data.split())
        if value:
            self.text.append(value)
            if self._current_href:
                self._current_anchor.append(value)


_MARKDOWN_LINK_RE = re.compile(r"\[[^\]]*\]\(([^)\s]+)(?:\s+[^)]*)?\)")


def extract_document_links(
    filename: str,
    content: bytes,
    *,
    base_url: str = "",
    source_system: str = "manual",
    source_version: str = "",
) -> list[DocumentLink]:
    """Return explicit http(s) document references without network access.

    HTML is the primary pilot surface. Markdown is included because its links
    are equally explicit and commonly arrives through SharePoint document
    libraries. Unsupported formats deliberately return no links rather than
    attempting extraction from rendered text or semantic similarity.
    """
    suffix = Path(filename).suffix.lower()
    raw = content.decode("utf-8", errors="replace")
    candidates: list[tuple[str, str, int]] = []
    if suffix in {".html", ".htm"}:
        parser = _HTMLTextAndLinks()
        parser.feed(raw)
        parser.close()
        candidates = parser.links
    elif suffix == ".md":
        candidates = [(href, "", index) for index, href in enumerate(_MARKDOWN_LINK_RE.findall(raw), start=1)]

    links: list[DocumentLink] = []
    seen: set[tuple[str, str, int]] = set()
    for href, anchor, index in candidates:
        target = urljoin(base_url, href).strip()
        try:
            link = DocumentLink(
                target_url=target,
                anchor_text=anchor,
                source_locator=f"anchor:{index}",
                source_system=source_system,
                source_version=source_version,
            )
        except ValueError:
            # Skip page fragments, mailto/javascript URLs and malformed values;
            # a document graph only accepts resolvable document identities.
            continue
        key = (link.target_url, link.anchor_text, index)
        if key not in seen:
            links.append(link)
            seen.add(key)
    return links


def load_structured_tables(
    filename: str,
    content: bytes,
    *,
    document_id: str,
    tenant: str = "default",
) -> list[StructuredTable]:
    """Return tables when a loader can extract them without inventing cells.

    Excel is already a structured source, so each worksheet becomes one table.
    PDF table extraction is intentionally not guessed from plain text; a future
    layout-aware extractor can supply the same ``structured_tables`` metadata
    contract and the ingestion writer will persist it unchanged.
    """
    if Path(filename).suffix.lower() != ".xlsx":
        return []
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    tables: list[StructuredTable] = []
    for index, sheet in enumerate(workbook.worksheets):
        values = [
            ["" if value is None else str(value) for value in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        while values and not any(values[-1]):
            values.pop()
        if not values:
            continue
        columns, *rows = values
        tables.append(StructuredTable(
            document_id=document_id,
            table_index=index,
            caption=sheet.title,
            columns=columns,
            rows=rows,
            extraction_method="xlsx_native",
            tenant=tenant,
        ))
    return tables
